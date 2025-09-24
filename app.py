# app.py - MVP Asistente Mercantil (FastAPI) - INGEST & DRAFT
import os, uuid, json, re
from typing import List, Optional
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import psycopg2, psycopg2.extras
from docx import Document

# Ingesta DOCX
import mammoth
from bs4 import BeautifulSoup

# --- Configuración por entorno ---
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
SUPABASE_CONN = os.getenv("SUPABASE_CONN", "")
ALLOWED_ORIGIN  = os.getenv("ALLOWED_ORIGIN", "*")
STORAGE_DIR     = os.getenv("STORAGE_DIR", "./data")

os.makedirs(STORAGE_DIR, exist_ok=True)

# --- App y CORS ---
app = FastAPI(title="IA Mercantil - Contestación a la Demanda (MVP)")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[ALLOWED_ORIGIN] if ALLOWED_ORIGIN != "*" else ["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Utilidades: DB ---
def db():
    if not SUPABASE_CONN:
        raise RuntimeError("Falta SUPABASE_CONN en variables de entorno")
    return psycopg2.connect(SUPABASE_CONN, sslmode="require")

# --- Utilidades: Embeddings y Chat ---
import requests

def embed_texts(texts: List[str]) -> List[List[float]]:
    """Embeddings con OpenAI (text-embedding-3-small)"""
    if not OPENAI_API_KEY:
        return [[0.0] * 1536 for _ in texts]  # stub para pruebas
    url = "https://api.openai.com/v1/embeddings"
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}"}
    out = []
    for t in texts:
        r = requests.post(url, headers=headers, json={"model":"text-embedding-3-small","input":t[:8000]})
        r.raise_for_status()
        out.append(r.json()["data"][0]["embedding"])
    return out

def chat_completion(messages: List[dict]) -> str:
    """Chat con OpenAI (gpt-4o-mini)"""
    if not OPENAI_API_KEY:
        return json.dumps({"formato_escrito": "Borrador (DUMMY). Añade OPENAI_API_KEY para salida real."}, ensure_ascii=False)
    url = "https://api.openai.com/v1/chat/completions"
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}"}
    payload = {"model":"gpt-4o-mini","temperature":0.2,"messages":messages}
    r = requests.post(url, headers=headers, json=payload)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]

# --- Parsing DOCX a secciones y párrafos ---
def docx_to_sections(docx_path: str):
    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(f)
    html = result.value
    soup = BeautifulSoup(html, "html.parser")

    sections = []
    current = {"id":"", "title":"", "paragraphs":[]}
    heading_counter = [0,0,0,0]  # h1..h4

    def commit_section():
        nonlocal current
        if current["title"] or current["paragraphs"]:
            sections.append(current)
        current = {"id":"", "title":"", "paragraphs":[]}

    for el in soup.find_all(["h1","h2","h3","h4","p","li","table"]):
        if el.name in ["h1","h2","h3","h4"]:
            level = int(el.name[1])
            heading_counter[level-1] += 1
            for i in range(level,4):
                heading_counter[i] = 0
            sec_num = ".".join([str(x) for x in heading_counter if x>0])
            commit_section()
            current["id"] = sec_num
            current["title"] = re.sub(r"\s+", " ", el.get_text(strip=True))
        else:
            txt = re.sub(r"\s+", " ", el.get_text(" ", strip=True))
            if txt:
                current["paragraphs"].append(txt)
    commit_section()
    return sections

def build_chunks_from_sections(sections, case_id):
    chunks = []
    for sec in sections:
        for i, p in enumerate(sec["paragraphs"], start=1):
            ref = f"Sección {sec['id'] or '0'} — {sec['title'] or 'Sin título'} — ¶{i}"
            chunks.append({
                "case_id": case_id,
                "section_id": sec["id"] or "0",
                "section_title": sec["title"] or "Sin título",
                "para_index": i,
                "ref": ref,
                "text": p
            })
    return chunks

# --- Endpoints ---
@app.get("/health")
def health():
    return {"ok": True}

@app.post("/ingest")
async def ingest(file: UploadFile = File(...), case_meta: str = Form("{}")):
    meta = json.loads(case_meta or "{}")
    case_id = str(uuid.uuid4())
    case_dir = os.path.join(STORAGE_DIR, case_id)
    os.makedirs(case_dir, exist_ok=True)

    # Guardar DOCX
    docx_path = os.path.join(case_dir, "original.docx")
    data = await file.read()
    with open(docx_path, "wb") as f:
        f.write(data)

    # Parsear a secciones/párrafos
    sections = docx_to_sections(docx_path)
    chunks = build_chunks_from_sections(sections, case_id)

    # Embeddings
    vectors = embed_texts([c["text"] for c in chunks])

    # Guardar en BD
    conn = db(); conn.autocommit = True
    with conn.cursor() as cur:
        cur.execute("insert into cases (id, meta) values (%s, %s)", (case_id, json.dumps(meta, ensure_ascii=False)))
        for c, v in zip(chunks, vectors):
            cur.execute("""
                insert into chunks (case_id, section_id, section_title, para_index, ref, text, embedding, tsv)
                values (%s, %s, %s, %s, %s, %s, %s, to_tsvector('spanish', %s))
            """, (case_id, c["section_id"], c["section_title"], c["para_index"], c["ref"], c["text"], v, c["text"]))
    conn.close()

    return {"case_id": case_id, "sections": len(sections), "chunks": len(chunks)}

class DraftIn(BaseModel):
    case_id: str
    rol: str = "demandado"
    jurisdiccion: str = "MX-mercantil"
    objetivo: str = "Contestar demanda mercantil"
    plazo: Optional[str] = None
    datos_expediente: dict
    preferencias: Optional[dict] = None

@app.post("/draft")
def draft(body: DraftIn):
    # Embedding de consulta (amplio)
    query_text = f"{body.objetivo}. {json.dumps(body.datos_expediente, ensure_ascii=False)}"
    qvec = embed_texts([query_text])[0]

    # BM25 terms (muy simple)
    terms = " & ".join([
        body.datos_expediente.get("actor",""),
        body.datos_expediente.get("demandado","")
    ]) or "demanda"

    # Retrieval híbrido
    conn = db()
    with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
        cur.execute("""
        SELECT id, section_id, section_title, para_index, ref, text,
          0.6*(1 - (embedding <=> %s)) + 0.4*ts_rank(tsv, to_tsquery('spanish', %s)) AS score
        FROM chunks
        WHERE case_id = %s
        ORDER BY score DESC
        LIMIT 12
        """, (qvec, terms, body.case_id))
        rows = cur.fetchall()
    conn.close()

    context = [f"[{r['ref']}] {r['text'][:1300]}" for r in rows]

    # PROMPT especializado en mercantil
    system = {
        "role": "system",
        "content": (
            "Actúas como abogado litigante especializado en derecho mercantil mexicano. "
            "Analiza la demanda con base en los fragmentos citados. "
            "No inventes hechos; si falta información, formula preguntas. "
            "Usa citas en formato (Sección S, ¶N). "
            "Devuelve JSON válido con claves: "
            "hechos_contestados, excepciones_y_defensas, argumentos_de_derecho, "
            "pruebas_ofrecidas, petitorios, riesgos_y_plazos, preguntas_pendientes, "
            "y formato_escrito (texto forense con encabezados y citas). "
            "Incluye checklist mercantil: competencia, personalidad, títulos/contratos, "
            "excepciones típicas (prescripción, pago, compensación), requisitos del CCom/LGTOC, pruebas."
        )
    }
    user = {
        "role": "user",
        "content": f"""
Rol: {body.rol}
Jurisdicción: {body.jurisdiccion}
Datos del expediente: {json.dumps(body.datos_expediente, ensure_ascii=False)}
Plazo: {body.plazo or ""}
Objetivo: {body.objetivo}

Fragmentos (citar como 'Sección S, ¶N'):
{"\n".join(context)}
"""
    }
    raw = chat_completion([system, user])

    try:
        data = json.loads(raw)
    except Exception:
        data = {"formato_escrito": raw, "preguntas_pendientes": ["Validar JSON devuelto por el modelo."]}

    # Generar DOCX
    case_dir = os.path.join(STORAGE_DIR, body.case_id)
    os.makedirs(case_dir, exist_ok=True)
    out_path = os.path.join(case_dir, "contestacion.docx")

    doc = Document()
    de = body.datos_expediente
    doc.add_paragraph(f"C. JUEZ DE LO MERCANTIL DE {de.get('ciudad','____________').upper()}")
    doc.add_paragraph(f"Expediente: {de.get('expediente','____________')}")
    doc.add_paragraph(f"Actor: {de.get('actor','____________')}   |   Demandado: {de.get('demandado','____________')}")
    doc.add_paragraph(" ")
    doc.add_heading("CONTESTACIÓN A LA DEMANDA (MERCANTIL)", level=1)
    doc.add_paragraph(data.get("formato_escrito",""))
    doc.add_page_break()
    doc.add_heading("PREGUNTAS PENDIENTES", level=2)
    for p in data.get("preguntas_pendientes", []):
        doc.add_paragraph(f"• {p}")
    doc.save(out_path)

    return {
        "summary": f"Borrador generado con {len(rows)} fragmentos citados.",
        "download_url": f"/files/cases/{body.case_id}/contestacion.docx",
        "preguntas_pendientes": data.get("preguntas_pendientes", [])
    }

@app.get("/files/cases/{case_id}/contestacion.docx")
def get_doc(case_id: str):
    path = os.path.join(STORAGE_DIR, case_id, "contestacion.docx")
    return FileResponse(path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="contestacion.docx")
